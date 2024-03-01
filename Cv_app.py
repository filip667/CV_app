import tkinter as tk
from tkinter import filedialog
import docx
from docx.shared import Inches
import re
import requests
from PIL import Image
from io import BytesIO

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def remove_sensitive_info(text):
    # Wyrażenia regularne do wykrywania i usuwania danych wrażliwych
    phone_regex = r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'
    email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    address_regex = r'\b\d{1,5}\s\w+\s\w+\b'
    cv_regex = r'\bCV\b|\bCurriculum Vitae\b'
    
    # Usunięcie danych wrażliwych
    text = re.sub(phone_regex, '', text)
    text = re.sub(email_regex, '', text)
    text = re.sub(address_regex, '', text)
    text = re.sub(cv_regex, '', text, flags=re.IGNORECASE)
    
    return text

def clear_cv():
    output_text.delete("1.0", tk.END)

def process_cv():
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        cv_text = extract_text_from_docx(file_path)
        sanitized_text = remove_sensitive_info(cv_text)
        output_text.delete("1.0", tk.END)
        output_text.insert(tk.END, sanitized_text)

def save_to_docx():
    sanitized_text = output_text.get("1.0", tk.END)
    doc = docx.Document()

    # Ustawienie czcionki na Times New Roman
    for style in doc.styles:
        if style.type == 1:
            style.font.name = 'Times New Roman'

    # Dodanie logo
    try:
        response = requests.get("https://upload.wikimedia.org/wikipedia/commons/thumb/2/24/LEGO_logo.svg/512px-LEGO_logo.svg.png")
        image = Image.open(BytesIO(response.content))
        image.save('A:/logo.png')
        doc.add_picture("logo.png", width=Inches(1.0))
    except Exception as e:
        print("Nie udało się dodać loga:", e)

    # Podzielenie tekstu na słowa i sprawdzenie, czy należą one do słów kluczowych
    for para_text in sanitized_text.split('\n'):
        p = doc.add_paragraph()
        for word in para_text.split():
            if any(word.lower() in keyword.lower() for keyword in ["doświadczenie", "work experience", "job experience", "experience", "języki", "languages", "osiągnięcia", "achievements", "osiągnięcia", "certyfikaty", "certificates", "umiejętności", "skills"]):
                run = p.add_run(word)
                run.bold = True  # Pogrubienie słowa kluczowego
            else:
                p.add_run(word)
            p.add_run(' ')

    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if save_path:
        doc.save(save_path)

# Utworzenie głównego okna
root = tk.Tk()
root.title("CV Sanitizer")

# Ramka dla pola tekstowego z CV
frame_cv = tk.Frame(root)
frame_cv.pack(fill=tk.BOTH, expand=True)

text_label = tk.Label(frame_cv, text="Wybierz plik CV:")
text_label.pack()

# Przeglądaj przycisk
button_browse = tk.Button(frame_cv, text="Przeglądaj", command=process_cv)
button_browse.pack()

# Ramka dla pola tekstowego z sformatowanym CV
frame_output = tk.Frame(root)
frame_output.pack(fill=tk.BOTH, expand=True)\

output_label = tk.Label(frame_output, text="Sformatowane CV:")
output_label.pack()

output_text = tk.Text(frame_output, height=10, width=30)
output_text.pack()

# Zmiana rozmiaru okna dla pola tekstowego
frame_output.config(width=500)

# Przyciski
button_clear = tk.Button(root, text="Wyczyść CV", command=clear_cv)
button_clear.pack()

button_save = tk.Button(root, text="Zapisz do DOCX", command=save_to_docx)
button_save.pack()

root.mainloop()
